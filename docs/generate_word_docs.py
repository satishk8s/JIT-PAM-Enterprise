#!/usr/bin/env python3
"""
Generate Word (.docx) documents from Markdown files.
Uses python-docx for document creation.
Run: python generate_word_docs.py
"""

import os
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILES = [
    '1_Current_State_Before.md',
    '2_Changes_Implemented.md',
    '3_Roadmap_Next_Steps.md',
]


def parse_md_to_structure(content):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        # Headings
        if line.startswith('# '):
            blocks.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            blocks.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            blocks.append(('h3', line[4:].strip()))
        elif line.startswith('#### '):
            blocks.append(('h4', line[5:].strip()))
        # Horizontal rule
        elif line.strip() == '---':
            blocks.append(('hr', None))
        # Table
        elif '|' in line and line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            blocks.append(('table', table_lines))
        # Code block
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append(('code', '\n'.join(code_lines)))
        # List item
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            blocks.append(('li', line.strip()[2:]))
        # Numbered list
        elif re.match(r'^\d+\.\s', line.strip()):
            blocks.append(('li', line.strip()))
        # Empty line
        elif not line.strip():
            blocks.append(('blank', None))
        # Paragraph
        else:
            blocks.append(('p', line.strip()))
        i += 1
    return blocks


def add_block(doc, block_type, content):
    """Add a block to the document."""
    if block_type == 'h1':
        p = doc.add_heading(content, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif block_type == 'h2':
        doc.add_heading(content, level=1)
    elif block_type == 'h3':
        doc.add_heading(content, level=2)
    elif block_type == 'h4':
        doc.add_heading(content, level=3)
    elif block_type == 'hr':
        doc.add_paragraph('_' * 60)
    elif block_type == 'p' and content:
        doc.add_paragraph(content)
    elif block_type == 'li':
        doc.add_paragraph(content, style='List Bullet')
    elif block_type == 'code' and content:
        p = doc.add_paragraph()
        p.style = 'Normal'
        run = p.add_run(content)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    elif block_type == 'table' and content:
        rows = []
        for r in content:
            if not r.strip():
                continue
            cells = r.strip().split('|')[1:-1]
            # Skip markdown table separator (|---|---|)
            if cells and all(re.match(r'^[\s\-:]+$', c.strip()) for c in cells):
                continue
            rows.append([c.strip() for c in cells])
        if rows:
            col_count = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = 'Table Grid'
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row[:col_count]):
                    table.rows[ri].cells[ci].text = cell.strip()
    elif block_type == 'blank':
        pass


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    blocks = parse_md_to_structure(content)
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for block_type, content in blocks:
        add_block(doc, block_type, content)
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")


def main():
    if not HAS_DOCX:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        return 1
    
    for md_file in MD_FILES:
        md_path = os.path.join(DOCS_DIR, md_file)
        if not os.path.exists(md_path):
            print(f"Warning: {md_file} not found")
            continue
        docx_name = md_file.replace('.md', '.docx')
        docx_path = os.path.join(DOCS_DIR, docx_name)
        convert_md_to_docx(md_path, docx_path)
    
    print("\nDone. Word documents created in docs/ folder.")
    return 0


if __name__ == '__main__':
    exit(main())

